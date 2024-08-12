import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX

csv_path = 'vis_dataset/mmlu_logicity_jessica/hard/test/hard_fixed_final_mmlu.csv'  # Update with the actual path to the CSV file
label_json_path = 'vis_dataset/mmlu_logicity_jessica/hard/test/hard_fixed_final_mmlu_label.json'  # Update with the actual path to the JSON file

def add_formatted_text(paragraph, text, color=None, bold=False):
    """
    Add text to a paragraph with specified formatting.
    """
    run = paragraph.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
    run.bold = bold

def matches_predicate_filter(self_predicates, predicates_filter):
    for filter_predicate in predicates_filter:
        if any(filter_predicate in sp for sp in self_predicates):
            return True
    return False

def get_data(num_questions, answer_filter, predicates_filter):
    with open(label_json_path, 'r') as f:
        labels = json.load(f)
    
    filtered_questions = []
    labels.shuffle()
    for label in labels:
        if len(filtered_questions) >= num_questions:
            break
        
        if answer_filter and label["answer"] != answer_filter:
            continue
        
        if predicates_filter and not matches_predicate_filter(label["self_predicates"], predicates_filter):
            continue
        
        filtered_questions.append(label)
    
    return filtered_questions

def generate_doc(filtered_questions, output_doc_path):
    doc = Document()
    doc.add_heading('Filtered Questions', 0)
    answer_map = {'A': 'Slow', 'B': 'Normal', 'C': 'Fast', 'D': 'Stop'}

    if not filtered_questions:
        print("No questions found with the given filters")
        return
    
    for q in filtered_questions:
        paragraph = doc.add_paragraph()
        
        question = q['question']
        answer = q['answer']
        self_entity = q['self_entity']
        self_predicates = q['self_predicates']
        predicates = [p.split('(')[0] for p in self_predicates]
        related_entity = q['related_entities']

        # Split the question into words or phrases
        words = question.split()

        for word in words:
            # Check if the word is part of any predicate or entity
            if any(p in word for p in predicates):
                if self_entity in word:
                    add_formatted_text(paragraph, word + ' ', color=(255, 0, 0), bold=True)
                elif any(e in word for e in related_entity):
                    add_formatted_text(paragraph, word + ' ', color=(0, 0, 255), bold=True)
                else:
                    add_formatted_text(paragraph, word + ' ', color=(0, 0, 0), bold=True)
            elif any(e in word for e in related_entity):
                add_formatted_text(paragraph, word + ' ', color=(0, 0, 255), bold=True)
            elif self_entity in word:
                add_formatted_text(paragraph, word + ' ', color=(255, 0, 0), bold=True)
            else:
                paragraph.add_run(word + ' ')

        doc.add_paragraph(f"Answer: {answer}, {answer_map[answer]}")
        doc.add_paragraph("")  # Add a blank line
    
    doc.save(output_doc_path)
    print(f"Document saved to {output_doc_path}")

if __name__ == "__main__":
    num_questions = 10  # Set how many questions you want
    answer_filter = 'C'  # Set the answer you want to filter by ('A', 'B', 'C', 'D'), or None for no filter
    # 'A': Slow, 'B': Normal, 'C': Fast, 'D': Stop
    predicates_filter = ['IsAmbulance', 'IsClose']  # Set the self-predicates you want to filter by, or an empty list for no filter
    # Unary: 'IsPedestrian', 'IsCar', 'IsAmbulance', 'IsBus', 'IsPolice', 'IsTiro', 'IsReckless', 'IsOld', 'IsYoung', 'IsAtInter', 'IsInInter'
    # Binary: 'IsClose', 'HigherPri', 'CollidingClose', 'LeftOf', 'RightOf', 'NextTo', 'Sees'

    # Closed world: All questioned Entities will be about: (One Of) 'IsAmbulance', 'IsBus', 'IsPolice', 'IsTiro', 'IsReckless'.
    output_doc_path = 'filtered_questions.docx'  # Set the path where you want to save the generated document
    filtered_questions = get_data(num_questions, answer_filter, predicates_filter)
    generate_doc(filtered_questions, output_doc_path)